import os
import pdfplumber
import docx
import re
import json
from pathlib import Path
from typing import List
import numpy as np
import praw

from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI

# keyword extraction
from keybert import KeyBERT

# Import our intelligent quality checker
from answer_quality_checker import AnswerQualityChecker

INDEX_PATH = "faiss_index"

# -----------------------------
# OpenRouter API Setup
# -----------------------------
os.environ["OPENROUTER_API_KEY"] = os.getenv("OPENROUTER_API_KEY")
os.environ["OPENAI_API_KEY"] = os.environ["OPENROUTER_API_KEY"]
os.environ["OPENAI_API_BASE"] = "https://openrouter.ai/api/v1"
os.environ["OPENAI_API_HEADERS"] = '{"HTTP-Referer":"https://huggingface.co", "X-Title":"GIKI-RAG-bot"}'

# -----------------------------
# Reddit Setup (OAuth)
# -----------------------------
reddit = praw.Reddit(
    client_id=os.getenv("REDDIT_CLIENT_ID"),
    client_secret=os.getenv("REDDIT_CLIENT_SECRET"),
    refresh_token=os.getenv("REDDIT_REFRESH_TOKEN"),
    user_agent=os.getenv("REDDIT_USER_AGENT")
)

# -----------------------------
# Keyword Extractor
# -----------------------------
kw_model = KeyBERT("all-MiniLM-L6-v2")

def extract_keywords(query: str, top_k=3):
    keywords = kw_model.extract_keywords(query, keyphrase_ngram_range=(1,2), stop_words='english', top_n=top_k)
    return [kw for kw, score in keywords]


def search_reddit_semantic(query: str, subreddit="giki", top_n=5, embeddings_model=None):
    """
    Fetch latest posts from r/giki and return top_n most semantically similar ones.
    Uses keyword extraction for better search.
    """
    # Step 1: Extract clean keywords/phrases
    keyphrases = extract_keywords(query, top_k=3)
    search_terms = keyphrases if keyphrases else [query]

    print(keyphrases)


    posts = []
    try:
        for term in search_terms:
            for submission in reddit.subreddit(subreddit).search(term, limit=20):
                posts.append({
                    "title": submission.title,
                    "selftext": submission.selftext,
                    "url": submission.url
                })
    except Exception:
        return []

    if not posts or embeddings_model is None:
        return []

    # Combine title + body
    post_texts = [f"{p['title']}\n{p['selftext']}" for p in posts]

    # Compute embeddings
    post_embeddings = embeddings_model.embed_documents(post_texts)
    query_embedding = embeddings_model.embed_query(query)

    # Cosine similarity
    similarities = [
        np.dot(query_embedding, p_emb) / (np.linalg.norm(query_embedding) * np.linalg.norm(p_emb))
        for p_emb in post_embeddings
    ]

    # Top N results
    top_indices = np.argsort(similarities)[::-1][:top_n]
    return [posts[i] for i in top_indices]


# -----------------------------
# Document Processor
# -----------------------------
class GIKIDocumentProcessor:
    def __init__(self, data_folder="data"):
        self.data_folder = data_folder
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        page_text = re.sub(r'\s+', ' ', page_text)
                        text += f"\n[Page {page_num + 1}]\n{page_text}\n"
        except Exception:
            pass
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
        except Exception:
            pass
        return text

    def extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception:
            return ""

    def extract_text_from_json(self, file_path: str) -> str:
        """Extract text from JSON files containing Reddit posts"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            text = ""
            if isinstance(data, list):
                for i, post in enumerate(data):
                    if isinstance(post, dict):
                        # Extract title and selftext
                        title = post.get('title', '')
                        selftext = post.get('selftext', '')
                        author = post.get('author', '')
                        post_id = post.get('id', '')
                        
                        # Combine post content
                        post_content = f"Title: {title}\nContent: {selftext}"
                        
                        # Add comments if available
                        comments = post.get('comments', [])
                        if comments:
                            post_content += "\nComments:\n"
                            for comment in comments:
                                if isinstance(comment, dict):
                                    comment_body = comment.get('body', '')
                                    comment_author = comment.get('author', '')
                                    if comment_body:
                                        post_content += f"- {comment_author}: {comment_body}\n"
                        
                        text += f"\n[Reddit Post {i+1} - ID: {post_id}]\n{post_content}\n"
            
            return text
        except Exception as e:
            # print(f"Error processing JSON file {file_path}: {str(e)}")
            return ""

    def load_documents(self) -> List[Document]:
        documents = []
        data_path = Path(self.data_folder)

        if not data_path.exists():
            return documents

        supported_extensions = {'.pdf', '.docx', '.txt', '.json'}

        for file_path in data_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                if file_path.suffix.lower() == '.pdf':
                    text = self.extract_text_from_pdf(str(file_path))
                elif file_path.suffix.lower() == '.docx':
                    text = self.extract_text_from_docx(str(file_path))
                elif file_path.suffix.lower() == '.txt':
                    text = self.extract_text_from_txt(str(file_path))
                elif file_path.suffix.lower() == '.json':
                    text = self.extract_text_from_json(str(file_path))
                else:
                    continue

                if text.strip():
                    chunks = self.text_splitter.split_text(text)

                    for i, chunk in enumerate(chunks):
                        if chunk.strip():
                            documents.append(
                                Document(
                                    page_content=chunk,
                                    metadata={
                                        "source": file_path.name,
                                        "chunk_id": i,
                                        "file_type": file_path.suffix.lower()
                                    }
                                )
                            )
        return documents


# -----------------------------
# bot
# -----------------------------
class GIKIbot:
    def __init__(self):
        self.qa_chain = None
        self.vectorstore = None
        self.processor = GIKIDocumentProcessor()
        self.quality_checker = AnswerQualityChecker()

        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

        self.prompt_template = """You are a helpful assistant for GIKI (Ghulam Ishaq Khan Institute of Engineering Sciences and Technology).
Answer questions based on official GIKI documents: prospectus, fee structure, academic rules, and handbook.

Instructions:
- Answer based only on context
- If answer not found, say "I don't have that information in the provided documents"
- Be specific and cite document source when possible
- Maintain professional, student-friendly tone

Context:
{context}

Question: {question}

Answer:"""

        self.custom_prompt = PromptTemplate(
            template=self.prompt_template,
            input_variables=["context", "question"]
        )

    def initialize_system(self):
        try:
            if os.path.exists(INDEX_PATH):
                self.vectorstore = FAISS.load_local(
                    INDEX_PATH,
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )
            else:
                documents = self.processor.load_documents()
                if not documents:
                    return "❌ No documents found. Add files to the 'data' folder."

                self.vectorstore = FAISS.from_documents(documents, self.embeddings)
                self.vectorstore.save_local(INDEX_PATH)

            # Initialize LLM
            llm = ChatOpenAI(
                model="deepseek/deepseek-r1-0528-qwen3-8b:free",
                base_url="https://openrouter.ai/api/v1",
                api_key=os.getenv("OPENAI_API_KEY"),
                default_headers={
                    "HTTP-Referer": "https://huggingface.co",
                    "X-Title": "GIKI-RAG-bot"
                },
                temperature=0.1
            )

            self.qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                ),
                return_source_documents=True,
                chain_type_kwargs={"prompt": self.custom_prompt}
            )

            return "✅ System ready! Ask questions now."
        except Exception as e:
            return f"❌ Error initializing system: {str(e)}"
    def ask_question(self, question: str) -> str:
        if not self.qa_chain:
            return "⚠️ System not initialized yet."

        if not question.strip():
            return "⚠️ Please enter a valid question."

        try:
            response = self.qa_chain.invoke({"query": question})
            answer = response["result"]
            source_docs = response["source_documents"]

            # Use intelligent quality checker to assess answer
            #print("🔍 Assessing answer quality...")
            quality_assessment = self.quality_checker.assess_answer_quality(question, answer)
            
            # Debug: Print quality assessment details
          #  print(f"🔍 Overall score: {quality_assessment['overall_score']:.2f}/10")
          #  print(f"🔍 AI assessment: {quality_assessment['ai_assessment']['ai_sufficient']}")
          #  print(f"🔍 AI reason: {quality_assessment['ai_assessment']['ai_reason']}")
          #  print(f"🔍 Basic scores: {quality_assessment['basic_scores']}")
           
            needs_fallback = not quality_assessment['is_sufficient']
            

           # print(f"🔍 Fallback triggered! Overall score: {quality_assessment['overall_score']:.2f}")
          # print(f"🔍 Reason: {quality_assessment['ai_assessment']['ai_reason']}")

            if needs_fallback:
                # print("📡 Document answer insufficient, searching Reddit...")
                top_posts = search_reddit_semantic(question, embeddings_model=self.embeddings)
                if top_posts:
                    # Build Reddit context
                    reddit_context = "\n\n".join([
                        f"**{p['title']}**\n{p['selftext'][:500]}...\n(Source: {p['url']})"
                        for p in top_posts if p['selftext']
                    ])
                    
                    # Get Reddit-based answer
                    llm = ChatOpenAI(
                        model="deepseek/deepseek-r1-0528-qwen3-8b:free",
                        base_url="https://openrouter.ai/api/v1",
                        api_key=os.getenv("OPENAI_API_KEY"),
                        temperature=0.2
                    )
                    reddit_answer = llm.invoke(
                        f"Answer the following question using the Reddit discussions:\n\n"
                        f"Question: {question}\n\nReddit Posts:\n{reddit_context}\n\nAnswer:"
                    ).content
                    
                    # Build Reddit sources
                    reddit_sources = set()
                    for post in top_posts:
                        reddit_sources.add(f"🌐 r/giki: {post['title'][:60]}...")
                    
                    reddit_source_text = "\n\nReddit Sources:\n" + "\n".join(reddit_sources) if reddit_sources else ""
                    
                    return f"⚠️ Not found in official documents. Based on Reddit discussions:\n\n{reddit_answer}{reddit_source_text}"
                else:
                    # No Reddit posts found either - return original answer with document sources
                    sources = set()
                    for doc in source_docs:
                        sources.add(f"📄 {doc.metadata['source']}")
                    source_text = "\n\nSources:\n" + "\n".join(sources) if sources else ""
                    return f"⚠️ {answer}\n\n(No additional information found on Reddit){source_text}"
            else:
                # Good document answer - return with document sources
                sources = set()
                for doc in source_docs:
                    sources.add(f"📄 {doc.metadata['source']}")
                source_text = "\n\nSources:\n" + "\n".join(sources) if sources else ""
                return f"{answer}{source_text}"
            
        except Exception as e:
            return f"❌ Error: {str(e)}"

